"""
Script to generate Phase 7 Word (.docx) documentation:
Phase_7_Live_Pipeline_UI_and_InApp_File_Viewer.docx
"""

import os
from pathlib import Path
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

OUTPUT_DIR = Path(__file__).resolve().parent.parent / "phase_documentation"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def set_cell_background(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)


def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def add_title(doc, title_text: str, subtitle_text: str):
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run = title_p.add_run(title_text)
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 58, 138)

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(18)
    sub_run = sub_p.add_run(subtitle_text)
    sub_run.font.size = Pt(12)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(100, 116, 139)


def add_heading_1(doc, text: str):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(16)
    h.paragraph_format.space_after = Pt(6)
    run = h.add_run(text)
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 64, 175)


def add_heading_2(doc, text: str):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(4)
    run = h.add_run(text)
    run.font.size = Pt(12.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor(51, 65, 85)


def add_bullet(doc, title: str, description: str):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(title + ": ")
    r1.bold = True
    r1.font.size = Pt(10.5)
    r2 = p.add_run(description)
    r2.font.size = Pt(10.5)


def add_callout(doc, text: str, alert_type="NOTE"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)

    if alert_type == "SUCCESS":
        bg_color = "ECFDF5"
        prefix = "✓ RESULT: "
    elif alert_type == "IMPORTANT":
        bg_color = "EFF6FF"
        prefix = "ℹ NOTE: "
    else:
        bg_color = "F8FAFC"
        prefix = "• "

    set_cell_background(cell, bg_color)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(prefix)
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(30, 58, 138)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def add_code_block(doc, code_text: str):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "0F172A")
    set_cell_margins(cell, top=140, bottom=140, left=160, right=160)

    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(226, 232, 240)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def generate_phase_7_doc():
    doc = docx.Document()

    add_title(
        doc,
        "Trace RAG — Phase 7: Live Pipeline UI & In-App File Viewer",
        "Server-Sent Events (SSE) Streaming, Live Progress Visualization, and In-App Modal Deep-Linking"
    )

    add_callout(
        doc,
        "Phase 7 delivers real-time transparency into the multi-stage RAG pipeline via Server-Sent Events (SSE), streaming execution updates live (route -> retrieve -> graph -> confidence -> answer -> verify -> done), and embeds an in-app viewer with deep-linking to exact document pages and media timestamps.",
        "IMPORTANT"
    )

    add_heading_1(doc, "1. Phase Objective & Core Capabilities")
    p = doc.add_paragraph()
    p.add_run(
        "The objective of Phase 7 is twofold: (1) wire the user interface directly to the real-time execution stages of the backend pipeline via "
        "Server-Sent Events without simulated delays, and (2) embed a rich, in-app file viewer supporting PDFs, Word (.docx) documents, images, audio, "
        "and video with automated page-seeking, timestamp deep-linking, and isolated page text extractions."
    )

    add_heading_1(doc, "2. Server-Sent Events (SSE) Streaming Architecture")
    add_heading_2(doc, "A. Backend SSE Endpoint (POST /api/query/stream)")
    add_bullet(doc, "Route Event (event: route)", "Emits the classified modality categories (e.g. document, image, conversational) and intent explanation.")
    add_bullet(doc, "Retrieve Event (event: retrieve)", "Emits the number of matched chunks and passage metadata following BM25 and vector retrieval.")
    add_bullet(doc, "Graph Event (event: graph)", "Emits multi-hop entity traversal paths and relation connections from the Knowledge Graph.")
    add_bullet(doc, "Confidence Event (event: confidence)", "Emits the Critic's confidence grade (HIGH/MEDIUM/LOW), diagnosis rationale, and missing aspects.")
    add_bullet(doc, "Retry Event (event: retry)", "Emits query reformulation details if the Critic triggers a self-correction retry.")
    add_bullet(doc, "Answer Event (event: answer)", "Emits the synthesized answer text containing [n] citation markers.")
    add_bullet(doc, "Verify Event (event: verify)", "Emits claim-by-claim groundedness scores and individual claim verification statuses.")
    add_bullet(doc, "Done Event (event: done)", "Emits the final complete SynthesisResult payload and triggers automatic message persistence.")

    sse_sample = (
        "event: route\n"
        "data: {\"stage\": \"route\", \"categories\": [\"document\", \"image\"], \"explanation\": \"Document inquiry\"}\n\n"
        "event: retrieve\n"
        "data: {\"stage\": \"retrieve\", \"chunks_count\": 5, \"chunks\": [...]}\n\n"
        "event: graph\n"
        "data: {\"stage\": \"graph\", \"hops_count\": 2, \"graph_hops\": [...]}\n\n"
        "event: confidence\n"
        "data: {\"stage\": \"confidence\", \"confidence\": \"high\", \"reason\": \"Strong factual support\"}\n\n"
        "event: answer\n"
        "data: {\"stage\": \"answer\", \"answer\": \"The LIMIT clause restricts rows [1]...\"}\n\n"
        "event: verify\n"
        "data: {\"stage\": \"verify\", \"citations\": [...], \"groundedness_score\": 1.0}\n\n"
        "event: done\n"
        "data: {\"stage\": \"done\", \"result\": {...}}"
    )
    add_code_block(doc, sse_sample)

    add_heading_2(doc, "B. Frontend Real-Time Stream Consumer (ChatSynthesisView.tsx)")
    add_bullet(doc, "ReadableStream Parser", "Consumes chunked UTF-8 byte streams asynchronously, parses SSE event/data blocks, and dispatches state updates.")
    add_bullet(doc, "Live Execution Tracker UI", "Renders glowing, responsive progress badges reflecting the active pipeline stage in real time.")

    add_heading_1(doc, "3. In-App File Viewer & Citation Deep-Linking")
    add_heading_2(doc, "A. Multimodal Inline Previewers (FileViewerModal.tsx)")
    add_bullet(doc, "PDF Documents", "Inline <iframe> with #page=N&toolbar=0 targeting to jump immediately to cited pages.")
    add_bullet(doc, "Word Documents (.docx)", "Native client-side Word document rendering via docx-preview with typography, tables, and margins preserved.")
    add_bullet(doc, "Audio & Video Media", "Native HTML5 <audio> and <video> elements with auto-seek to citation timestamps (e.g. 01:24).")
    add_bullet(doc, "Isolated Page Extraction", "Displays only the extracted text for the cited page (e.g. [Page 8]), with evidence statement highlighting and full-doc toggle.")

    add_heading_1(doc, "4. Verification & Automated Test Results")
    add_callout(
        doc,
        "Automated integration test suite backend.tests.test_phase7_streaming verified both conversational and document RAG streaming flows over SSE, confirming sequential event ordering and payload completeness.",
        "SUCCESS"
    )

    test_output = (
        "Ran 2 tests in 18.412s\n"
        "test_01_conversational_stream ... OK\n"
        "test_02_document_rag_stream ... OK\n\n"
        "OK"
    )
    add_code_block(doc, test_output)

    output_file = OUTPUT_DIR / "Phase_7_Live_Pipeline_UI_and_InApp_File_Viewer.docx"
    doc.save(str(output_file))
    print(f"Generated: {output_file}")


if __name__ == "__main__":
    generate_phase_7_doc()
