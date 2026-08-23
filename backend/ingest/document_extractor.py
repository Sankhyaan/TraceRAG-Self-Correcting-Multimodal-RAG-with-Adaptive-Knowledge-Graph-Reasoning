import io
import re
from pathlib import Path
from typing import Optional
import fitz  # PyMuPDF
from PIL import Image
import pytesseract
from backend.config import get_settings

settings = get_settings()

# Configure Tesseract path if specified in config
if settings.tesseract_cmd:
    pytesseract.pytesseract.tesseract_cmd = settings.tesseract_cmd


class DocumentExtractor:
    """Extracts text from PDF documents, Word files, and plain text with OCR fallback."""

    def extract(self, filename: str, file_bytes: bytes) -> str:
        ext = Path(filename).suffix.lower()

        if ext == ".pdf":
            return self._extract_pdf(file_bytes)
        elif ext in (".docx", ".doc"):
            return self._extract_docx(file_bytes)
        elif ext in (".txt", ".md", ".csv", ".json", ".log"):
            return self._extract_text(file_bytes)
        else:
            # Fallback to general text decode
            return self._extract_text(file_bytes)

    def _clean_text(self, text: str) -> str:
        """Sanitizes PDF/document text by removing control characters, fixing broken bullets, and unprintable glyphs."""
        if not text:
            return ""
        # Remove control characters (except newline and tab)
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", " ", text)
        # Replace missing glyphs / private-use area characters with bullet or space
        # e.g., \ue000-\uf8ff, \ufff0-\uffff, \U00010000-\U0010ffff, \uf0a7, \uf0b7, \uf0d8, \uf0e0-\uf0ff, \ufffd
        text = re.sub(r"[\ue000-\uf8ff\ufff0-\uffff\U00010000-\U0010ffff]", " ", text)
        text = re.sub(r"[􀀀🗎\uf0a7\uf0b7\uf0d8\uf0e0-\uf0ff]", "- ", text)
        # Clean stray box or repeated symbol characters at line starts
        text = re.sub(r"^[ \t]*[□■◆◇○●•▪▫-][ \t]*", "- ", text, flags=re.MULTILINE)
        # Collapse multiple empty lines
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def _extract_pdf(self, pdf_bytes: bytes) -> str:
        """
        Extracts text from each page of a PDF using PyMuPDF.
        If a page contains little or no text (< 40 characters), falls back to OCR via Pytesseract.
        """
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        total_pages = len(doc)
        pages_output = []

        for page_num in range(total_pages):
            page = doc.load_page(page_num)
            text = self._clean_text(page.get_text())

            # Check if page has sufficient extractable digital text
            alphanumeric_count = len(re.findall(r"[a-zA-Z0-9]", text))

            if alphanumeric_count >= 40:
                # Direct digital text
                pages_output.append(f"[Page {page_num + 1}]\n{text}")
            else:
                # Scanned or image-based page -> Run OCR
                ocr_text = self._clean_text(self._ocr_pdf_page(page))
                if ocr_text.strip():
                    pages_output.append(f"[Page {page_num + 1} (OCR)]\n{ocr_text.strip()}")
                elif text:
                    pages_output.append(f"[Page {page_num + 1}]\n{text}")
                else:
                    pages_output.append(f"[Page {page_num + 1}]\n(No readable text found on page)")

        doc.close()
        return "\n\n".join(pages_output)

    def _ocr_pdf_page(self, page: fitz.Page) -> str:
        """Renders a PDF page to high-res image and runs Tesseract OCR."""
        try:
            # Render at 150 DPI for fast and accurate OCR
            pix = page.get_pixmap(dpi=150)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            ocr_text = pytesseract.image_to_string(img)
            return ocr_text.strip()
        except Exception as e:
            # If Tesseract binary is not installed or errors, log and return empty
            print(f"[DocumentExtractor OCR Notice] Pytesseract execution notice: {str(e)}")
            return ""

    def _extract_docx(self, docx_bytes: bytes) -> str:
        """Extracts text from Microsoft Word documents (.docx) including tables."""
        try:
            import docx

            doc = docx.Document(io.BytesIO(docx_bytes))
            sections = []

            # Extract paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    sections.append(text)

            # Extract tables
            for t_idx, table in enumerate(doc.tables):
                table_rows = []
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells]
                    if any(row_cells):
                        table_rows.append(" | ".join(row_cells))
                if table_rows:
                    sections.append(f"\n[Table {t_idx + 1}]\n" + "\n".join(table_rows))

            return "\n\n".join(sections) if sections else "(Empty Word Document)"
        except Exception as e:
            raise RuntimeError(f"Word document extraction failed: {str(e)}")

    def _extract_text(self, text_bytes: bytes) -> str:
        """Extracts plain text with UTF-8 / latin-1 fallback."""
        try:
            return text_bytes.decode("utf-8")
        except UnicodeDecodeError:
            return text_bytes.decode("latin-1", errors="replace")


document_extractor = DocumentExtractor()
